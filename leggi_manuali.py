from docx import Document
import os, sys
sys.stdout.reconfigure(encoding='utf-8')

paths = [
    r"Z:\Progetti IT\2026\Portale\Manuali-Utente\Manuale-Utente-Agenzie-Pallet-Hostess.docx",
    r"Z:\Progetti IT\2026\Portale\Manuali-Utente\Manuale-Utente-Welfare.docx",
]

for path in paths:
    print("=" * 70)
    print(os.path.basename(path))
    print("=" * 70)
    doc = Document(path)
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        try:
            style = p.style.name if p.style else "Normal"
        except Exception:
            style = "Normal"
        print(f"[{style}] {p.text}")
    for i, tbl in enumerate(doc.tables):
        print(f"--- TABELLA {i+1} ---")
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                print(" | ".join(cells))
    print()
