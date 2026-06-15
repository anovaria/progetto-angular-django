"""
Genera Manuale-Utente-CaricaPromo.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = r"Z:\Progetti IT\2026\Portale\Manuali-Utente\Manuale-Utente-CaricaPromo.docx"

# ── colori brand ──
BLU       = RGBColor(0x1A, 0x52, 0x76)   # reparto
VIOLA     = RGBColor(0x6C, 0x34, 0x83)   # abbig
VERDE     = RGBColor(0x1E, 0x84, 0x49)
ARANCIO   = RGBColor(0xD3, 0x54, 0x00)
GRIGIO_BG = RGBColor(0xF5, 0xF5, 0xF5)
BIANCO    = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get('val', 'single'))
        tag.set(qn('w:sz'), kwargs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'DDDDDD'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


doc = Document()

# ── imposta margini ──
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── stili base ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(10)

# ────────────────────────────────────────────────────────────────
# FRONTESPIZIO
# ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('PORTALE INTRANET — GROS CIDAC')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.bold = False

doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('Manuale Utente')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = BLU

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('Carico Promozioni')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = BLU

doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('Freschi / Reparto  ·  Abbigliamento')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run(f'Versione 1.0  —  {datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()


# ────────────────────────────────────────────────────────────────
# helpers
# ────────────────────────────────────────────────────────────────
def h1(text, color=BLU):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = color
    p.runs[0].font.size = Pt(16)
    return p

def h2(text, color=BLU):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = color
    p.runs[0].font.size = Pt(13)
    return p

def h3(text, color=None):
    p = doc.add_heading(text, level=3)
    if color:
        p.runs[0].font.color.rgb = color
    p.runs[0].font.size = Pt(11)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10) if p.runs else None
    return p

def nota(text, color=None):
    p = doc.add_paragraph()
    run = p.add_run('  ℹ  ' + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = color or RGBColor(0x55, 0x55, 0x55)
    return p

def attenzione(text):
    p = doc.add_paragraph()
    run = p.add_run('  ⚠  ' + text)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part and bold_part in text:
        idx = text.index(bold_part)
        if idx > 0:
            p.add_run(text[:idx]).font.size = Pt(10)
        r = p.add_run(bold_part)
        r.bold = True
        r.font.size = Pt(10)
        after = text[idx+len(bold_part):]
        if after:
            p.add_run(after).font.size = Pt(10)
    else:
        p.add_run(text).font.size = Pt(10)
    return p

def tabella_sconti():
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Tipo Sconto', 'Quando usarlo', 'Campi abilitati', 'Esempio']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = BIANCO
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(hdr[i], '1A5276')

    rows_data = [
        ('PE', 'Sconto percentuale\n(anche composto)', 'Sconto Extra +\nSconto 1 + Sconto 2', 'Es. 5% extra\no 5%+2%+1%'),
        ('E',  'Sconto in euro fisso', 'Solo campo "Sconto €"', 'Es. € 0,50 a pezzo'),
        ('OM', 'Omaggio (X paga Y)', 'Qta Acquisto +\nQta Omaggio', 'Es. acquista 6,\npaga 5'),
    ]
    for rd in rows_data:
        row = tbl.add_row().cells
        for i, val in enumerate(rd):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_border(row[i])
    doc.add_paragraph()

def tabella_differenze():
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'

    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Funzionalità', 'Freschi / Reparto', 'Abbigliamento']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = BIANCO
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(hdr[i], '1A5276')

    rows_data = [
        ('Caricamento da CCOM',          '✔', '✔'),
        ('Caricamento manuale (barcode)', '✔', '✗'),
        ('Duplica da promo storica Gold', '✔', '✗'),
        ('Importa da file Excel',         '✗', '✔'),
        ('Logica blocco sconti per tipo', '✗ (tutti liberi)', '✔ (si sbloccano\nin base al tipo)'),
        ('Abilita Prezzo Vendita',        '✗',  '✔ (checkbox APV)'),
        ('Visualizza inseriti ordinabile','✔ (colonne cliccabili)', '✗'),
        ('Colonne Sell-in nel visualizza','✔', '✗'),
    ]
    for rd in rows_data:
        row = tbl.add_row().cells
        for i, val in enumerate(rd):
            row[i].text = val
            run = row[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if i > 0:
                run.font.color.rgb = VERDE if val.startswith('✔') else (RGBColor(0xC0,0x39,0x2B) if val.startswith('✗') else RGBColor(0,0,0))
            set_cell_border(row[i])
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════
# 1. INTRODUZIONE
# ════════════════════════════════════════════════════════════════
h1('1. Introduzione')
body(
    'Il modulo Carico Promozioni consente di preparare e accodare le promozioni per i prodotti '
    'Freschi/Reparto e Abbigliamento. Le due versioni condividono lo stesso flusso operativo '
    'di base ma offrono funzionalità di caricamento articoli diverse.'
)
doc.add_paragraph()

h2('Flusso generale (uguale per entrambe le versioni)')
for step in [
    ('1', 'Seleziona il CCOM (Contratto Commerciale) del fornitore.'),
    ('2', 'Carica gli articoli disponibili per quel fornitore.'),
    ('3', 'Seleziona gli articoli da includere nella promo.'),
    ('4', 'Imposta i parametri della promozione (piano, date, sconto).'),
    ('5', 'Clicca "Inserisci la Promo" per accodare gli articoli all\'export.'),
    ('6', 'Verifica il riepilogo con "Visualizza Articoli Inseriti".'),
]:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(f'  {step[1]}')
    r.font.size = Pt(10)

doc.add_paragraph()
nota('Gli articoli accodati rimangono in memoria fino all\'esportazione finale dal modulo Scarico Promo.')
doc.add_paragraph()

h2('Differenze tra le due versioni')
tabella_differenze()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 2. CARICO PROMO FRESCHI / REPARTO
# ════════════════════════════════════════════════════════════════
h1('2. Carico Promo Freschi / Reparto', color=BLU)
nota('Accessibile dal menu laterale: Promozioni → Carico Promo Freschi')
doc.add_paragraph()

# ── 2.1 ──
h2('2.1  Selezione CCOM e caricamento articoli')
body('Nella sezione 1 della pagina:')
bullet('Cliccare nel campo CCOM e digitare il codice o il nome del fornitore per filtrare l\'elenco.', 'CCOM')
bullet('Selezionare il fornitore dall\'elenco a tendina.')
bullet('Cliccare il pulsante Carica Articoli.', 'Carica Articoli')
body('Appare la griglia con tutti gli articoli disponibili per quel fornitore.')
doc.add_paragraph()
nota('Se sono presenti articoli già inseriti in precedenza, all\'apertura della pagina appare un avviso arancione in alto a destra.')
doc.add_paragraph()

# ── 2.2 ──
h2('2.2  Metodi alternativi di caricamento articoli')

h3('Caricamento Manuale (barcode)', color=ARANCIO)
body(
    'Permette di aggiungere singoli articoli tramite lettore barcode o digitazione manuale, '
    'indipendentemente dal CCOM selezionato.'
)
bullet('Cliccare il pulsante Caricamento Manuale.', 'Caricamento Manuale')
bullet('Scansionare il barcode oppure digitare il codice articolo e premere Invio o Cerca.')
bullet('Ripetere per ogni articolo da aggiungere. Gli articoli già presenti nella lista vengono segnalati.')
bullet('Cliccare Conferma e Aggiungi per trasferire gli articoli nella griglia principale, già selezionati.', 'Conferma e Aggiungi')
bullet('Il pulsante Svuota lista permette di ricominciare la scansione.', 'Svuota lista')
doc.add_paragraph()

h3('Duplica da Promo Gold', color=BLU)
body(
    'Permette di copiare gli articoli da una promozione storica già presente nel gestionale Gold, '
    'utile per rinnovare promozioni ricorrenti.'
)
bullet('Cliccare il pulsante Duplica da Promo.', 'Duplica da Promo')
bullet('Selezionare la promozione storica dall\'elenco (mostra codice, descrizione piano e date).')
bullet('Selezionare il fornitore (CCOM) oppure scegliere "★ Tutti i fornitori" per copiare tutti.')
bullet('Verificare l\'anteprima degli articoli trovati.')
bullet('Cliccare Carica Articoli per inserirli nella griglia.', 'Carica Articoli')
doc.add_paragraph()

# ── 2.3 ──
h2('2.3  Selezione articoli')
body('Una volta caricati gli articoli, utilizzare la barra di selezione per scegliere quelli da includere nella promo:')
bullet('Click su una riga per selezionarla/deselezionarla (diventa verde se selezionata).')
bullet('Seleziona tutti / Deseleziona — seleziona o deseleziona l\'intera lista.', 'Seleziona tutti')
bullet('Per linea prodotto — seleziona tutti gli articoli di una linea.', 'Per linea prodotto')
bullet('Per tipo riordino — seleziona per tipo di riordino.', 'Per tipo riordino')
bullet('Per fascia prezzo — seleziona per fascia di prezzo.', 'Per fascia prezzo')
body('Il contatore "Selezionati: N" si aggiorna in tempo reale.')
doc.add_paragraph()
nota('La ricerca articolo in alto a destra permette di trovare e scorrere rapidamente a un articolo specifico nella lista.')
doc.add_paragraph()

# ── 2.4 ──
h2('2.4  Parametri promozione')

h3('Piano Promo')
bullet('Usare il Filtro Piano per restringere l\'elenco per lettera.', 'Filtro Piano')
bullet('Selezionare il Piano Promo dall\'elenco: le date Sell-out e Sell-in si compilano automaticamente.', 'Piano Promo')
bullet('Le date possono essere modificate manualmente se necessario.')
doc.add_paragraph()

h3('Date')
bullet('Sell-out Inizio / Fine: date in cui la promo è visibile al consumatore (facoltative ma consigliate).', 'Sell-out')
bullet('Sell-in Inizio / Fine: date di ordine al fornitore (OBBLIGATORIE).', 'Sell-in')
attenzione('Se le date Sell-out non sono valorizzate, il sistema chiede conferma prima di procedere.')
doc.add_paragraph()

h3('Meccanica')
bullet('Selezionare la meccanica promozionale dal menu apposito (es. sconto volantino, esposizione, ecc.).')
doc.add_paragraph()

h3('Tipo Sconto e importi')
body('Selezionare il Tipo Sconto e compilare i campi corrispondenti:')
tabella_sconti()

body('Il campo Tot. Sconto mostra lo sconto percentuale composto calcolato automaticamente.')
doc.add_paragraph()

# ── 2.5 ──
h2('2.5  Inserimento e verifica')
bullet('Cliccare Inserisci la Promo per accodare gli articoli selezionati con i parametri inseriti.', 'Inserisci la Promo')
bullet('Se vengono rilevati duplicati (articoli già presenti nell\'export), appare un avviso: si può accodare comunque o annullare.')
bullet('Il contatore "In export: N" in basso a destra si aggiorna.')
bullet('Cliccare Visualizza Articoli Inseriti per controllare il riepilogo completo.', 'Visualizza Articoli Inseriti')
bullet('Dal riepilogo è possibile Stampare o Svuotare tutto l\'export.', 'Stampare')
doc.add_paragraph()

# ── 2.6 ──
h2('2.6  Pulisci schermo')
body(
    'Il pulsante Pulisci Schermo elimina gli articoli dalla griglia di lavoro (fase 1) '
    'senza intaccare quelli già accodati nell\'export. Utilizzarlo per cambiare fornitore '
    'o ricominciare la selezione.'
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 3. CARICO PROMO ABBIGLIAMENTO
# ════════════════════════════════════════════════════════════════
h1('3. Carico Promo Abbigliamento', color=VIOLA)
nota('Accessibile dal menu laterale: Promozioni → Carico Promo Abbigliamento')
doc.add_paragraph()

body(
    'Il flusso è identico a quello della versione Freschi/Reparto (sezione 2), '
    'con le differenze descritte di seguito.'
)
doc.add_paragraph()

# ── 3.1 ──
h2('3.1  Importa da file Excel', color=VIOLA)
body(
    'Invece del caricamento manuale e del duplica da promo, '
    'la versione Abbigliamento permette di importare articoli da un file Excel fornito dal fornitore.'
)
bullet('Cliccare il pulsante Importa Excel.', 'Importa Excel')
bullet('Selezionare il file .xlsx o .xls.')
bullet('Cliccare Importa per caricare gli articoli nella griglia.', 'Importa')
doc.add_paragraph()
attenzione(
    'Il file Excel deve contenere le colonne: CODART, DESCRART, CODFORN, CCOM, '
    'DESCRCCOM, pv_std, PVOFF ARR. File con struttura diversa non verranno importati correttamente.'
)
doc.add_paragraph()

# ── 3.2 ──
h2('3.2  Logica Tipo Sconto (specifica Abbigliamento)', color=VIOLA)
body(
    'A differenza della versione Reparto, nella versione Abbigliamento i campi sconto '
    'si abilitano o disabilitano automaticamente in base al Tipo Sconto selezionato:'
)
tabella_sconti()

body('I campi disabilitati appaiono in grigio e non sono modificabili.')
doc.add_paragraph()

h3('Abilita Prezzo Vendita (checkbox APV)')
body(
    'Per impostare anche il prezzo di vendita promozionale, spuntare la checkbox '
    '"Abilita Prz. Vendita": si attivano i campi Meccanica e Valore / Prezzo.'
)
doc.add_paragraph()
nota('Nella versione Reparto i campi Meccanica e sconti sono sempre editabili senza checkbox.')
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 4. MESSAGGI E SITUAZIONI PARTICOLARI
# ════════════════════════════════════════════════════════════════
h1('4. Messaggi e situazioni particolari')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for i, h in enumerate(['Messaggio / Situazione', 'Causa', 'Soluzione']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.color.rgb = BIANCO
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr[i], '1A5276')

situazioni = [
    (
        'Avviso arancione "Hai N articoli già inseriti" all\'apertura',
        'Sessione precedente non completata',
        'Cliccare "Articoli Inseriti" per verificare cosa è rimasto in coda prima di continuare.'
    ),
    (
        'Date Sell-in obbligatorie',
        'Il campo Sell-in Inizio o Fine è vuoto',
        'Compilare entrambe le date Sell-in prima di cliccare "Inserisci la Promo".'
    ),
    (
        'Avviso duplicati: "N articoli già presenti"',
        'Gli articoli sono già presenti nell\'export con lo stesso codice',
        'Scegliere "Accoda comunque" per aggiungere una seconda riga, oppure "Annulla" per non duplicare.'
    ),
    (
        'Nessun articolo selezionato',
        'Non è stato selezionato alcun articolo nella griglia',
        'Selezionare almeno un articolo (click sulla riga o "Seleziona tutti") prima di procedere.'
    ),
    (
        'Articolo non trovato (caricamento barcode)',
        'Il codice scansionato non è presente in Gold',
        'Verificare il codice articolo o il barcode sul prodotto.'
    ),
    (
        '"Articolo già nella lista" (caricamento barcode)',
        'Il codice è stato scansionato due volte',
        'Continuare con la scansione: il duplicato è stato ignorato automaticamente.'
    ),
    (
        'Avviso "Date Sell-out non valorizzate"',
        'I campi Sell-out sono vuoti',
        'Cliccare "Procedi senza Sell-out" per continuare, oppure "Annulla" per inserire le date.'
    ),
]
for s in situazioni:
    row = tbl.add_row().cells
    for i, val in enumerate(s):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_border(row[i])

doc.add_paragraph()
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 5. GLOSSARIO
# ════════════════════════════════════════════════════════════════
h1('5. Glossario')

termini = [
    ('CCOM', 'Contratto Commerciale — identifica il rapporto commerciale con un fornitore specifico.'),
    ('Sell-out', 'Periodo in cui la promo è attiva per il consumatore finale in negozio.'),
    ('Sell-in', 'Periodo in cui il negozio ordina la merce al fornitore a condizioni promozionali.'),
    ('Piano Promo', 'Campagna promozionale di riferimento (es. volantino, evento stagionale).'),
    ('Sconto Extra (PE)', 'Sconto percentuale aggiuntivo applicato al prezzo base, componibile con Sconto 1 e Sconto 2.'),
    ('Tot. Sconto', 'Sconto percentuale totale calcolato come composto di Sconto Extra, Sconto 1 e Sconto 2.'),
    ('Omaggio (OM)', 'Meccanica "acquista N paga M": il cliente riceve unità gratuite al raggiungimento di una soglia.'),
    ('Fase 1', 'Griglia di lavoro temporanea per la selezione articoli, visibile solo all\'utente corrente.'),
    ('Export / Accodamento', 'Insieme di articoli pronti per essere esportati nel gestionale tramite il modulo Scarico Promo.'),
    ('Duplica da Promo', 'Funzione (solo Reparto) per copiare gli articoli di una promozione storica da Gold.'),
    ('VL', 'Codice punto vendita / livello di listino associato all\'articolo.'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for i, h in enumerate(['Termine', 'Definizione']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.color.rgb = BIANCO
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr[i], '1A5276')

for term, defn in termini:
    row = tbl.add_row().cells
    row[0].text = term
    row[0].paragraphs[0].runs[0].bold = True
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].text = defn
    row[1].paragraphs[0].runs[0].font.size = Pt(9)
    for c in row:
        set_cell_border(c)

doc.add_paragraph()

# ── footer ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'GROS CIDAC s.r.l.  —  Portale Intranet  —  {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── salva ──
doc.save(OUTPUT)
print(f'Salvato: {OUTPUT}')
